import os
import time
import urllib.request
from docx import Document
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.common.exceptions import StaleElementReferenceException as StaleException, ElementNotVisibleException, \
    ElementNotSelectableException, TimeoutException
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC

# complex = set of courses

chrome_options = Options()
chrome_options.add_experimental_option('prefs', {'download.prompt_for_download': 'false',
                                                 'download.directory_upgrade': 'true',
                                                 'safebrowsing.enabled': 'false',
                                                 'profile.default_content_setting_values.automatic_downloads': 1,
                                                 'profile.default_content_settings.popups': 0
                                                 })
chrome_options.add_argument('--mute-audio')
chrome_driver = r'E:/PyCharmProjects/CourseraComplexesSceletons/chromedriver.exe'
driver = webdriver.Chrome(executable_path=chrome_driver, options=chrome_options)
driver.implicitly_wait(20)

wait = WebDriverWait(driver, 20, poll_frequency=1, ignored_exceptions=[StaleException,
                                                                       ElementNotSelectableException,
                                                                       ElementNotVisibleException,
                                                                       ])

root_folder = r'D:/Coursera'


def video_downloading(file_link, file_folder, file_name, week_unsorted_folder):
    """
    function download video and additional materials from video page
    """

    href = file_link.find_element_by_xpath('../../../../..').get_attribute('href')
    driver.execute_script("window.open('');")
    driver.switch_to.window(driver.window_handles[2])
    driver.get(href)

    time.sleep(4.5)
    drop_down_list_xpath = '//button[@id="downloads-dropdown-btn"]'
    drop_down_list = wait.until(EC.element_to_be_clickable((By.XPATH, drop_down_list_xpath)))
    drop_down_list.click()

    webvtt_sub = driver.find_element_by_xpath('//li[@class="menuitem"][2]/a').get_attribute('href')
    txt_sub = driver.find_element_by_xpath('//li[@class="menuitem"][3]/a').get_attribute('href')
    urllib.request.urlretrieve(webvtt_sub, file_folder + f'\\subtitles-en.vtt')
    urllib.request.urlretrieve(txt_sub, file_folder + f'\\subtitle.txt')
    video_file_link = driver.find_element_by_xpath('//video').get_attribute('src')

    driver.get(video_file_link)
    # time.sleep(2)
    video_xpath = '//video/source'
    wait.until(EC.presence_of_element_located((By.XPATH, video_xpath)))
    video = driver.find_element_by_xpath(video_xpath).get_attribute('src')
    try:
        urllib.request.urlretrieve(video, file_folder + f'\\{file_name}.webm')
        print('Video was downloaded.')
    except FileNotFoundError:
        urllib.request.urlretrieve(video, week_unsorted_folder + f'\\{file_name}.webm')
        print('Video was downloaded, but not sorted.')

    driver.close()
    driver.switch_to.window(driver.window_handles[1])


def reading_downloading(file_link, file_and_path, file_name, week_unsorted_folder):
    """
    function copy reading's text and paste it in the specific .docx file
    """
    href = file_link.find_element_by_xpath('../../../../..').get_attribute('href')
    driver.execute_script("window.open('');")
    driver.switch_to.window(driver.window_handles[2])
    driver.get(href)

    doc_header_xpath = '//div[@class="reading-title"]/h2'
    doc_content_xpath = '//div[@class="rc-CML styled"]'
    time.sleep(3.8)

    try:
        doc_header = wait.until(EC.visibility_of_element_located((By.XPATH, doc_header_xpath))).text
    except TimeoutException:
        doc_header = "Reading"
    doc_content = wait.until(EC.visibility_of_element_located((By.XPATH, doc_content_xpath))).text

    document = Document()
    document.add_heading(doc_header, 0)
    document.add_paragraph(doc_content)

    try:
        document.save(file_and_path)
    except FileNotFoundError:
        document.save(week_unsorted_folder + f'\\{file_name}.docx')

    driver.close()
    driver.switch_to.window(driver.window_handles[1])


def course_directories_tree_deploying(course_folder, main_tab):
    """
    function browse&serf all the materials of the course
    and makes directories on your computer
    """
    driver.switch_to.window(driver.window_handles[1])
    weeks_list_xpath = '//div[@class="rc-NavigationDrawer"]/a'
    time.sleep(4.5)
    wait.until(EC.presence_of_all_elements_located((By.XPATH, weeks_list_xpath)))
    weeks_list = wait.until(EC.visibility_of_all_elements_located((By.XPATH, weeks_list_xpath)))
    week_number = 0

    for week in weeks_list:
        video_files_increment = 0
        reading_files_increment = 0
        week_number += 1
        time.sleep(3.5)
        week.click()
        time.sleep(3.5)
        week_title = wait.until(EC.visibility_of_element_located((By.XPATH, '//section/div[1]/div/h3'))).text

        # template for .translate() method
        table = week_title.maketrans('', '', r'\/:*?"<>|')

        week_title = week_title.translate(table)

        files_xpath = '//a[@class="nostyle"]/div/div/div/div/div[@class="rc-WeekItemName headline-1-text"]'
        files = wait.until(EC.visibility_of_all_elements_located((By.XPATH, files_xpath)))

        week_folder = os.path.join(course_folder, f'Week {week_number} - {week_title}')
        week_videos_folder = os.path.join(week_folder, '(1)Videos')
        week_readings_folder = os.path.join(week_folder, '(2)Readings')
        week_unsorted_folder = os.path.join(week_folder, 'Unsorted')

        if not os.path.exists(week_folder):
            os.makedirs(week_videos_folder)
            os.makedirs(week_readings_folder)
            os.makedirs(week_unsorted_folder)

        # cycle for videos' and readings' links
        for file_link in files:

            file_name = file_link.text
            file_name = file_name.translate(table)
            file_name = file_name.replace('\n', '')

            if file_name[0:5] == 'Video':
                junk = ['Fun', 'Demonstration', 'Video', '']
                for i in junk:
                    file_name = file_name.replace('VideoLecture' + i, '', 1)
                video_files_increment += 1
                file_folder = os.path.join(week_videos_folder, f'({video_files_increment}){file_name}')
                if not os.path.exists(file_folder):
                    os.makedirs(file_folder)
                video_downloading(file_link, file_folder, file_name, week_unsorted_folder)

            if file_name[0:7] == 'Reading':
                file_name = file_name.replace('Reading ', '', 1)
                reading_files_increment += 1
                file_and_path = os.path.join(week_readings_folder, f'({reading_files_increment}){file_name}.docx')
                reading_downloading(file_link, file_and_path, file_name, week_unsorted_folder)

    driver.close()
    driver.switch_to.window(main_tab)
    print('Course Tab MUST BE CLOSED NOW')


def main():
    login = input('Enter your Coursera login:\n')
    password = input('Enter your Coursera password:\n')
    driver.get(
        'https://www.coursera.org/programs/kyiv-national-university-of-trade-and-economics-ya-atc4j?authMode=login')

    main_tab = driver.window_handles[0]

    login_textbox = wait.until(EC.presence_of_element_located((By.XPATH, '//input[@type="email"]')))
    password_textbox = wait.until(EC.presence_of_element_located((By.XPATH, '//input[@type="password"]')))
    submit_button = wait.until(EC.presence_of_element_located((By.XPATH, '//button[@data-js="submit"]')))

    login_textbox.click()
    login_textbox.send_keys(login)
    password_textbox.click()
    password_textbox.send_keys(password)
    submit_button.click()

    input('Press "Enter" WHEN captcha will be resolved:')
    print('*skeleton - template of structured directories and files')
    complex_title = input('Type in title of the complex, which skeleton* you want to create (copy-paste from site):\n')
    complex_block = f'//h3[text()="{complex_title}"]/../..'

    courses_list = wait.until(EC.presence_of_all_elements_located((By.XPATH, f'{complex_block}/div[2]/div/div/ul/li')))
    complex_folder = os.path.join(root_folder, f'{complex_title}')

    # Creating main folder for complex/specialization
    if not os.path.exists(complex_folder):
        os.mkdir(complex_folder)

    course_number_increment = 0
    for course_number in courses_list:
        course_number_increment += 1
        course_number.click()
        time.sleep(3.5)
        course_title_xpath = f'{complex_block}/div[2]/div/div[2]/div/div/div/div/div/div/div[2]/div/div[1]'
        course_title = wait.until(EC.presence_of_element_located((By.XPATH, course_title_xpath))).text
        course_title = course_title.replace(':', '')
        course_folder = os.path.join(complex_folder, f'({course_number_increment}){course_title}')
        course_preview_button_xpath = f'{complex_block}//div[2]/div[2]//a'
        course_preview_button = wait.until(EC.element_to_be_clickable((By.XPATH, course_preview_button_xpath)))
        course_preview_button.click()
        course_directories_tree_deploying(course_folder, main_tab)
    print(f'Complex "{complex_title}" was fully downloaded.')


if __name__ == '__main__':
    main()
